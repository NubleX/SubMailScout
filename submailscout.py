import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urljoin, urlparse, parse_qs, urlencode
import time
import random
from fake_useragent import UserAgent
import logging
import aiohttp
import asyncio
from typing import Set, Tuple, Dict, List
import tqdm
import dns.resolver
import os
import mimetypes
import fitz  # PyMuPDF for PDF handling
from docx import Document  # python-docx for DOCX handling
import xlrd  # for Excel files
import magic  # python-magic for file type detection
from yarl import URL
import hashlib
import json
import urllib.parse
import io

class WebScanner:
    def __init__(self, base_url: str, max_workers: int = 10, timeout: int = 30):
        self.base_url = base_url if base_url.startswith(('http://', 'https://')) else f'http://{base_url}'
        self.domain = urlparse(self.base_url).netloc
        self.max_workers = max_workers
        self.timeout = timeout
        self.session = self._create_session()
        self.visited_urls = set()
        self.visited_files = set()
        self.rate_limiter = asyncio.Semaphore(5)
        self.download_dir = 'downloaded_files'
        os.makedirs(self.download_dir, exist_ok=True)
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('scanner.log'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def _create_session(self) -> aiohttp.ClientSession:
        return aiohttp.ClientSession(
            headers={
                "User-Agent": UserAgent().random,
                "Accept": "*/*",
                "Accept-Language": "en-US,en;q=0.5",
                "Accept-Encoding": "gzip, deflate",
                "Connection": "keep-alive",
            },
            timeout=aiohttp.ClientTimeout(total=self.timeout)
        )

    def _is_same_domain(self, url: str) -> bool:
        """Check if URL belongs to the target domain."""
        try:
            return urlparse(url).netloc.endswith(self.domain)
        except:
            return False

    def _normalize_url(self, url: str) -> str:
        """Normalize URL to avoid duplicates."""
        try:
            parsed = URL(url)
            # Sort query parameters
            if parsed.query_string:
                params = parse_qs(parsed.query_string)
                sorted_params = urlencode(sorted(params.items()), doseq=True)
                return str(parsed.with_query(sorted_params))
            return str(parsed)
        except:
            return url

    async def fetch_url(self, url: str, is_file: bool = False) -> Tuple[bytes, str, int]:
        """Fetch URL content with enhanced file type detection."""
        async with self.rate_limiter:
            try:
                if not url.startswith(('http://', 'https://')):
                    url = urljoin(self.base_url, url)
                
                url = self._normalize_url(url)
                self.logger.info(f"Fetching: {url}")
                await asyncio.sleep(random.uniform(0.5, 1.5))
                
                async with self.session.get(url, ssl=False) as response:
                    content = await response.read()
                    content_type = response.headers.get('Content-Type', '')
                    return content, content_type, response.status
            except Exception as e:
                self.logger.error(f"Error fetching {url}: {str(e)}")
                return b"", "", 0

    def find_emails(self, content: str) -> Set[str]:
        """Extract email addresses with improved pattern matching."""
        email_pattern = r'''[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'''
        emails = set(re.findall(email_pattern, content, re.IGNORECASE))
        return {email for email in emails if self._is_valid_email(email)}

    def extract_emails_from_file(self, file_path: str) -> Set[str]:
        """Extract emails from various file types."""
        emails = set()
        try:
            mime = magic.Magic(mime=True)
            file_type = mime.from_file(file_path)
            
            if 'pdf' in file_type:
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                emails.update(self.find_emails(text))
                
            elif 'document' in file_type:
                doc = Document(file_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                emails.update(self.find_emails(text))
                
            elif 'excel' in file_type or 'spreadsheet' in file_type:
                workbook = xlrd.open_workbook(file_path)
                for sheet in workbook.sheets():
                    for row in range(sheet.nrows):
                        for col in range(sheet.ncols):
                            cell_value = str(sheet.cell_value(row, col))
                            emails.update(self.find_emails(cell_value))
                            
        except Exception as e:
            self.logger.error(f"Error extracting emails from {file_path}: {str(e)}")
            
        return emails

    async def close(self):
        """Close the aiohttp session."""
        if self.session:
            await self.session.close()

    async def scan(self) -> Dict:
        self.logger.info(f"Starting enhanced scan of {self.base_url}")
        start_time = time.time()

        all_emails = set()
        directories = await self.scan_directories()
        subdomains = await self.enumerate_subdomains()

        # Gather all URLs to scan
        urls_to_scan = {self.base_url}
        urls_to_scan.update(directories)
        urls_to_scan.update(f"http://{sub}" for sub in subdomains)

        # Recursively scan each URL
        for url in urls_to_scan:
            try:
                emails = await self.recursive_scan(url)
                all_emails.update(emails)
            except Exception as e:
                self.logger.error(f"Error scanning URL {url}: {str(e)}")

        elapsed_time = time.time() - start_time
        
        results = {
            "emails": sorted(list(all_emails)),
            "directories": sorted(list(directories)),
            "subdomains": sorted(list(subdomains)),
            "scan_time": f"{elapsed_time:.2f} seconds",
            "total_urls_scanned": len(self.visited_urls),
            "total_files_processed": len(self.visited_files)
        }

        with open('scan_results.json', 'w') as f:
            json.dump(results, f, indent=4)

        self.logger.info(f"Scan completed in {elapsed_time:.2f} seconds")
        return results

    def _is_processable_url(self, url: str) -> bool:
        """Check if URL potentially contains processable content."""
        # File extensions to look for
        extensions = {
            '.xlsx', '.xls', '.doc', '.docx', '.pdf',
            '.txt', '.csv', '.rtf', '.xml', '.json'
        }
    
        # Convert URL to lowercase for case-insensitive matching
        url_lower = url.lower()
    
        # Check file extensions
        if any(ext in url_lower for ext in extensions):
            return True
    
        # Check common document paths
        doc_paths = {
            '/documents/', '/docs/', '/files/',
            '/download/', '/downloads/', '/shared/',
            '/media/', '/uploads/', '/resource',
            '/assets/', '/public/', '/publications/'
        }
    
        if any(path in url_lower for path in doc_paths):
            return True
    
        return False  # Ensure return is at the correct indentation
    
    async def find_documents(self, content: str, base_url: str) -> Set[str]:
            """Find document links in content."""
            documents = set()
            try:
                soup = BeautifulSoup(content, 'html.parser')
                
                # Find all links
                for a in soup.find_all('a', href=True):
                    href = a['href']
                    # Handle both relative and absolute URLs
                    full_url = urljoin(base_url, href)
                    
                    # Decode URL-encoded characters
                    decoded_url = urllib.parse.unquote(full_url)
                    
                    if self._is_same_domain(decoded_url) and self._is_processable_url(decoded_url):
                        documents.add(decoded_url)
                        
                # Also look for embedded frames and objects
                for elem in soup.find_all(['iframe', 'embed', 'object']):
                    src = elem.get('src', elem.get('data', ''))
                    if src:
                        full_url = urljoin(base_url, src)
                        decoded_url = urllib.parse.unquote(full_url)
                        if self._is_same_domain(decoded_url) and self._is_processable_url(decoded_url):
                            documents.add(decoded_url)
                        
            except Exception as e:
                self.logger.error(f"Error finding documents: {str(e)}")
                
            return documents
    
    async def process_file_content(self, content: bytes, content_type: str) -> Set[str]:
        """Process file content safely in memory."""
        emails = set()
        try:
            # Handle Excel files
            if 'spreadsheet' in content_type or content_type.endswith(('xlsx', 'xls')):
                # Create in-memory file-like object
                import io
                xlsx_file = io.BytesIO(content)
                wb = xlrd.open_workbook(file_contents=content)
                for sheet in wb.sheets():
                    for row in range(sheet.nrows):
                        for col in range(sheet.ncols):
                            cell_value = str(sheet.cell_value(row, col))
                            emails.update(self.find_emails(cell_value))
                            
            # Handle PDF files
            elif 'pdf' in content_type:
                with io.BytesIO(content) as pdf_stream:
                    pdf = fitz.open(stream=pdf_stream, filetype="pdf")
                    text = ""
                    for page in pdf:
                        text += page.get_text()
                    emails.update(self.find_emails(text))
                    
            # Handle Word documents
            elif 'word' in content_type or content_type.endswith(('doc', 'docx')):
                with io.BytesIO(content) as doc_stream:
                    doc = Document(doc_stream)
                    text = ""
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    emails.update(self.find_emails(text))
                    
        except Exception as e:
            self.logger.error(f"Error processing file content: {str(e)}")
            
        return emails
    
    async def recursive_scan(self, url: str, depth: int = 3) -> Set[str]:
        """Recursively scan URLs for content and files."""
        if depth <= 0 or url in self.visited_urls:
            return set()
            
        self.visited_urls.add(url)
        emails = set()
        
        try:
            content, content_type, status = await self.fetch_url(url)
            
            if status != 200:
                return emails
                
            # Process regular web pages
            if 'text/html' in content_type:
                # Extract emails from HTML content
                html_content = content.decode('utf-8', errors='ignore')
                emails.update(self.find_emails(html_content))
                
                # Find document links
                doc_urls = await self.find_documents(html_content, url)
                # Find dynamic pages
                dynamic_urls = await self.find_dynamic_pages(html_content, url)
                
                # Combine all discovered URLs
                all_urls = doc_urls.union(dynamic_urls)
                
                # Process discovered URLs
                for discovered_url in all_urls:
                    if discovered_url not in self.visited_urls:
                        sub_emails = await self.recursive_scan(discovered_url, depth - 1)
                        emails.update(sub_emails)
                        
            # Process documents directly
            elif self._is_processable_url(url):
                file_emails = await self.process_file_content(content, content_type)
                emails.update(file_emails)
                
        except Exception as e:
            self.logger.error(f"Error in recursive scan of {url}: {str(e)}")
            
        return emails
    
        def _is_valid_email(self, email: str) -> bool:
            if email.endswith(('.js', '.css', '.jpg', '.png', '.gif', '.svg')):
                return False
            pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
            return bool(re.match(pattern, email))
    
        async def download_and_process_file(self, url: str, content_type: str) -> Set[str]:
            """Download and process files for email extraction."""
            emails = set()
            try:
                content, _, status = await self.fetch_url(url, is_file=True)
                if status == 200 and content:
                    file_hash = hashlib.md5(content).hexdigest()
                    if file_hash not in self.visited_files:
                        self.visited_files.add(file_hash)
                        
                        ext = mimetypes.guess_extension(content_type) or '.bin'
                        file_path = os.path.join(self.download_dir, f"file_{file_hash}{ext}")
                        
                        with open(file_path, 'wb') as f:
                            f.write(content)
                        
                        emails.update(self.extract_emails_from_file(file_path))
                        
                        # Clean up file after processing
                        try:
                            os.remove(file_path)
                        except:
                            pass
                            
            except Exception as e:
                self.logger.error(f"Error processing file {url}: {str(e)}")
                
            return emails
    
        async def find_dynamic_pages(self, content: str, base_url: str) -> Set[str]:
            """Find dynamic pages and pagination links."""
            urls = set()
            try:
                soup = BeautifulSoup(content, 'html.parser')
                
                # Find all links
                for a in soup.find_all('a', href=True):
                    href = a['href']
                    full_url = urljoin(base_url, href)
                    
                    # Check if it's a dynamic page
                    if self._is_same_domain(full_url):
                        parsed = urlparse(full_url)
                        if parsed.query or '.php' in parsed.path:
                            urls.add(self._normalize_url(full_url))
                            
                # Look for pagination patterns
                patterns = [
                    r'page=\d+',
                    r'p=\d+',
                    r'offset=\d+',
                    r'start=\d+',
                    r'board=\d+\.\d+',
                    r'topic=\d+\.\d+'
                ]
                
                for pattern in patterns:
                    matches = re.finditer(pattern, content)
                    for match in matches:
                        # Generate next/prev page URLs
                        num = re.search(r'\d+', match.group())
                        if num:
                            num = int(num.group())
                            for i in range(max(0, num-1), num+2):
                                new_url = re.sub(r'\d+', str(i), full_url)
                                if self._is_same_domain(new_url):
                                    urls.add(self._normalize_url(new_url))
                                    
            except Exception as e:
                self.logger.error(f"Error finding dynamic pages: {str(e)}")
                
            return urls
    
        async def recursive_scan(self, url: str, depth: int = 3) -> Set[str]:
            """Recursively scan URLs for content and files."""
            if depth <= 0 or url in self.visited_urls:
                return set()
                
            self.visited_urls.add(url)
            emails = set()
            
            try:
                content, content_type, status = await self.fetch_url(url)
                
                if status != 200:
                    return emails
                    
                # Handle different content types
                if 'text/html' in content_type:
                    # Extract emails from HTML content
                    emails.update(self.find_emails(content.decode('utf-8', errors='ignore')))
                    
                    # Find and scan dynamic pages
                    dynamic_urls = await self.find_dynamic_pages(content.decode('utf-8', errors='ignore'), url)
                    for dynamic_url in dynamic_urls:
                        if dynamic_url not in self.visited_urls:
                            sub_emails = await self.recursive_scan(dynamic_url, depth - 1)
                            emails.update(sub_emails)
                            
                # Handle downloadable files
                elif any(file_type in content_type.lower() for file_type in 
                        ['pdf', 'msword', 'spreadsheet', 'excel', 'document']):
                    file_emails = await self.download_and_process_file(url, content_type)
                    emails.update(file_emails)
                    
            except Exception as e:
                self.logger.error(f"Error in recursive scan of {url}: {str(e)}")
                
            return emails
    
        async def scan_directories(self) -> Set[str]:
            common_paths = [
                "admin", "login", "dashboard", "user", "api", "wp-admin", 
                "uploads", "images", "includes", "js", "css", "static",
                "media", "download", "downloads", "content", "assets",
                "backup", "db", "sql", "dev", "test", "staging"
            ]
            
    async def scan_directories(self) -> Set[str]:
            common_paths = [
                "admin", "login", "dashboard", "user", "api", "wp-admin", 
                "uploads", "images", "includes", "js", "css", "static",
                "media", "download", "downloads", "content", "assets",
                "backup", "db", "sql", "dev", "test", "staging"
            ]
            
            directories = set()
            async def check_directory(path: str):
                try:
                    url = urljoin(self.base_url, path)
                    self.logger.info(f"Checking directory: {url}")
                    content, content_type, status = await self.fetch_url(url)  # Fixed unpacking
                    
                    if status == 200:
                        directories.add(url)
                        for ext in ['.php', '.txt', '.html', '.xml', '.json']:
                            file_url = urljoin(url + '/', 'index' + ext)
                            file_content, file_type, file_status = await self.fetch_url(file_url)  # Fixed unpacking
                            if file_status == 200:
                                directories.add(file_url)
                except Exception as e:
                    self.logger.error(f"Error checking directory {path}: {str(e)}")
    
            tasks = [check_directory(path) for path in common_paths]
            await asyncio.gather(*tasks)
            return directories
    
    async def enumerate_subdomains(self) -> Set[str]:
            subdomains = set()
            
            async def check_dns_record(subdomain: str):
                try:
                    answers = await asyncio.get_event_loop().run_in_executor(
                        None, dns.resolver.resolve, f"{subdomain}.{self.domain}", 'A'
                    )
                    if answers:
                        subdomains.add(f"{subdomain}.{self.domain}")
                except:
                    pass
    
            common_subdomains = ['www', 'mail', 'remote', 'blog', 'webmail', 'server',
                               'ns1', 'ns2', 'smtp', 'secure', 'vpn', 'api', 'dev',
                               'staging', 'test', 'portal', 'admin']
    
            tasks = [check_dns_record(sub) for sub in common_subdomains]
            await asyncio.gather(*tasks)
            
            try:
                url = f"https://crt.sh/?q=%.{self.domain}&output=json"
                async with self.session.get(url) as response:
                    if response.status == 200:
                        data = await response.json()
                        for entry in data:
                            name = entry.get('name_value', '').lower()
                            if name.endswith(self.domain):
                                subdomains.add(name)
            except Exception as e:
                self.logger.error(f"Error querying certificate logs: {e}")
    
            return subdomains
    
            async def close(self):
                await self.session.close()

async def main():
    print("Enter the target domain (e.g., example.com):", end=" ")
    target_domain = input().strip()
    
    scanner = WebScanner(target_domain)
    try:
        results = await scanner.scan()
        
        print("\n=== Scan Results ===")
        print(f"\nEmails found ({len(results['emails'])}):")
        for email in results['emails']:
            print(f"  - {email}")
            
        print(f"\nDirectories found ({len(results['directories'])}):")
        for directory in results['directories']:
            print(f"  - {directory}")
            
        print(f"\nSubdomains found ({len(results['subdomains'])}):")
        for subdomain in results['subdomains']:
            print(f"  - {subdomain}")
            
        print(f"\nScan completed in {results['scan_time']}")
        print(f"Total URLs scanned: {results['total_urls_scanned']}")
        print("\nFull results have been saved to 'scan_results.json'")
        
    finally:
        await scanner.close()

if __name__ == "__main__":
    asyncio.run(main())
